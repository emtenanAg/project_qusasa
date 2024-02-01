from django.shortcuts import render, redirect
from django.contrib.auth import login, logout
from .forms import CustomUserCreationForm  # Import at the top
from django.contrib.auth import authenticate, login
from django.http import HttpResponse
from .models import CustomUser
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
# Create your views here.
from django.shortcuts import render
from django.contrib.auth.models import User
from django.http import HttpResponseRedirect
from django.urls import reverse
from .utils import searchByQuery, extractIdFromUrl, analyse_channels, video_analysis, analyze_playlist, analyze_channel, download_audio_from_youtube, transcribe_youtube_video, summarize_youtube_video, create_word_document, topic_analysis, get_realted_videos
from .youtube_api import get_youtube_client
from django.core.exceptions import ValidationError
import re
import os
from django.conf import settings
from decouple import config
from .models import TopicAnalysisHistory
import openai
import pandas as pd
def custom_admin(request):
    users = User.objects.all()
    context = {
        'users': users,
    }
    return render(request, 'admin/custom_admin.html', context)


def email_verified_required(function):
    def wrap(request, *args, **kwargs):
        if request.user.is_verified:
            return function(request, *args, **kwargs)
        else:
            return redirect('confirm_email')
    return wrap

def home(request):
    return render(request, 'qusasa/home.html')

@login_required
@email_verified_required
def base(request):
    topic_histories = TopicAnalysisHistory.objects.filter(user=request.user).order_by('-created_at')
    video_histories = VideoAnalysisHistory.objects.filter(user=request.user).order_by('-created_at')
    playlist_histories = PlaylistAnalysisHistory.objects.filter(user=request.user).order_by('-created_at')
    channel_histories = ChannelAnalysisHistory.objects.filter(user=request.user).order_by('-created_at')
    video_retrieving_histories = VideoRetrievingHistory.objects.filter(user=request.user).order_by('-created_at')
    competitive_histories = CompetitiveAnalysisHistory.objects.filter(user=request.user).order_by('-created_at')

    return render(request, 'qusasa/base.html', {
        'topic_histories': topic_histories,
        'video_histories': video_histories,
        'playlist_histories': playlist_histories,
        'channel_histories': channel_histories,
        'video_retrieving_histories': video_retrieving_histories,
        'competitive_histories': competitive_histories,
    })

def signup(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            user.backend = 'qusasa.backends.EmailBackend'
            
            return redirect('login')
    else:
        form = CustomUserCreationForm()
    return render(request, 'qusasa/signup.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']
        user = authenticate(request, email=email, password=password)
        if user is not None:
            login(request, user)
            if user.is_verified:
                return redirect('base')
            else:
                # Redirect to a page where they need to input the confirmation code
                return redirect('confirm_email')  # Change to the name/url of your confirmation code input page
        else:
            return render(request, 'qusasa/login.html', {'error': 'Invalid email or password.'})
    return render(request, 'qusasa/login.html')

def logout_view(request):
    logout(request)
    return redirect('login')


@staff_member_required
@email_verified_required
def admin_only_pages(request):
    return render(request,'login.html')

def YouTubeFeat(request):
    return render(request,'qusasa/YouTubeFeat.html')
def InstagramFeat(request):
    return render(request,'qusasa/InstagramFeat.html')
def wFeature(request):
    return render(request,'qusasa/wFeature.html')


@login_required
def confirm_email(request):
    if request.method == 'POST':
        input_code = request.POST['confirmation_code']
        user = request.user
        if user.email_confirmation_code == input_code:
            user.is_verified = True
            user.email_confirmation_code = None  # Clear the confirmation code
            user.save()
            return redirect('base')
        else:
            # Handle incorrect code, maybe show an error message on the confirmation page
            return render(request, 'qusasa/confirm_email.html', {'error': 'Invalid confirmation code.'})
    else:
        # Display the page where they input the confirmation code
        return render(request, 'qusasa/confirm_email.html')
 
@staff_member_required
def inquiries_view(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'qusasa/inquiries.html')

@login_required
def competitive_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'features_pages/competitive_analysis/competitive_analysis_details.html')

from django.http import HttpResponseRedirect
from formtools.wizard.views import SessionWizardView
from .forms import CompetitiveAnalysisTypeForm, myChannelPlaylistInputForm, YouTubeSearchForm, YouTubeCategorySearchForm, ChannelsListInput, FindInitialChoiceForm
from .models import CompetitiveAnalysisHistory


class CompetitiveAnalysisWizard(SessionWizardView):
    form_list = [CompetitiveAnalysisTypeForm, myChannelPlaylistInputForm, FindInitialChoiceForm, ChannelsListInput]
    template_name = 'features_pages/competitive_analysis/competitive_analysis.html'
    
    
    def get_form_initial(self, step):
        initial = super().get_form_initial(step)
        history_id = self.kwargs.get('history_id')

        if history_id:
            history = get_object_or_404(CompetitiveAnalysisHistory, id=history_id, user=self.request.user)
            if step == '0':
                initial.update({
                    'analysis_type': history.analysis_type,
                    'input_text': history.input_text,
                })
            elif step == '2':
                initial.update({
                    'choice': history.choice,
                })
            elif step == '3':
                if history.choice == 'search':
                    initial.update({
                        'search_query': history.search_query,
                        'order': history.order,
                        'region_code': history.region_code,
                        'language': history.language,
                    })
                elif history.choice == 'input_list':
                    # Assuming you store multiple channel URLs in a JSON field
                    channel_urls = history.channel_urls
                    for i, url in enumerate(channel_urls):
                        initial.update({f'channel_url_{i+1}': url})
        return initial
    

    
    def get_form(self, step=None, data=None, files=None):
        form = super().get_form(step, data, files)
        # If it's the step that needs conditional logic to pick the form.
        if step == '3':
            previous_choice = self.get_cleaned_data_for_step('2')['choice']  # '1' is the step for FindInitialChoiceForm
            if previous_choice == 'input_list':
                self.form_list[step] = ChannelsListInput 
                form = ChannelsListInput(data, files, prefix=self.get_form_prefix(step, ChannelsListInput))
            elif previous_choice == 'search':
                self.form_list[step] = YouTubeSearchForm 
                form = YouTubeSearchForm(data, files, prefix=self.get_form_prefix(step, YouTubeSearchForm))
            elif previous_choice == 'category':
                self.form_list[step] = YouTubeCategorySearchForm 
                form = YouTubeCategorySearchForm(data, files, prefix=self.get_form_prefix(step, YouTubeCategorySearchForm))
        return form
    
    def get_form_kwargs(self, step=None):
        kwargs = super().get_form_kwargs(step)
        if step in ['1', '2', '3']:  # Include all steps after the initial selection
            analysis_type = self.request.session.get('analysis_type')
            if analysis_type:
                kwargs['analysis_type'] = analysis_type
            else:
                # Redirect to step 0 if analysis_type is missing
                self.storage.current_step = '0'
                return self.render_goto_step('0')
        return kwargs
    
    def post(self, *args, **kwargs):
        if self.steps.current == '0' and '0-analysis_type' in self.request.POST:
            # Store analysis type in the session after completing step 0
            analysis_type = self.request.POST['0-analysis_type']
            self.request.session['analysis_type'] = analysis_type
        return super().post(*args, **kwargs)

    def get_context_data(self, form, **kwargs):
        context = super().get_context_data(form=form, **kwargs)
        # Determine the current step. If it's the second form, change the label dynamically
        if self.steps.current == '1':  # assuming '1' is the index of InputForm
            previous_data = self.get_cleaned_data_for_step('0') or {}
            analysis_type = previous_data.get('analysis_type')
            if analysis_type == 'channel':
                form.fields['input_text'].label = 'channel'
            elif analysis_type == 'playlist':
                form.fields['input_text'].label = 'playlist'
        
        if self.steps.current == '3':
            context['step2_data'] = self.get_cleaned_data_for_step('2') or {}
        return context
    
    
    def done(self, form_list, **kwargs):
        # Process the cleaned data
        cleaned_data = self.get_all_cleaned_data()
        channel_urls = []
        # Do something with the data and generate dataset/report as required
        entity_type = cleaned_data.get('analysis_type')
        my_link = cleaned_data.get('input_text')
        search_or_list = cleaned_data.get('choice') #('input_list', 'search')
    
        
        youtube = get_youtube_client()
        
        if(search_or_list == 'search'):
            query = cleaned_data.get('search_query')
            order = cleaned_data.get('order')
            region_code = cleaned_data.get('region_code')
            language = cleaned_data.get('language')
            ids_list = searchByQuery(youtube, query, entity_type, order, region_code, language)
            ids_list.insert(0, extractIdFromUrl(my_link))
            
            
            
        elif(search_or_list == 'input_list'):
            ids_list = [extractIdFromUrl(my_link)]
        for i in range(1, 5):
            url_field = f'channel_url_{i}'
            if url_field in cleaned_data and cleaned_data[url_field]:
                channel_url = extractIdFromUrl(cleaned_data[url_field])
                ids_list.append(channel_url)
                channel_urls.append(channel_url)
        
        history_id = self.kwargs.get('history_id')

        if history_id:
            # Update existing history record
            history = get_object_or_404(CompetitiveAnalysisHistory, id=history_id, user=self.request.user)
            history.analysis_type = cleaned_data.get('analysis_type')
            history.input_text = cleaned_data.get('input_text')
            history.choice = cleaned_data.get('choice')
            history.search_query = cleaned_data.get('search_query', '')
            history.order = cleaned_data.get('order', '')
            history.region_code = cleaned_data.get('region_code', '')
            history.language = cleaned_data.get('language', '')
            history.channel_urls = channel_urls
            history.save()
        else:
            # Create new history record
            CompetitiveAnalysisHistory.objects.create(
                user=self.request.user,
                analysis_type=cleaned_data.get('analysis_type'),
                input_text=cleaned_data.get('input_text'),
                choice=cleaned_data.get('choice'),
                search_query=cleaned_data.get('search_query', ''),
                order=cleaned_data.get('order', ''),
                region_code=cleaned_data.get('region_code', ''),
                language=cleaned_data.get('language', ''),
                channel_urls=channel_urls,
            )
            
        channel_data_df, top_videos_df, channel_icons, durations_list = analyse_channels(ids_list, entity_type, youtube)
        if 'engagementScore' in top_videos_df.columns:
            top_videos_df.drop('engagementScore', axis=1, inplace=True)
        channel_data_csv = channel_data_df.to_csv(index=False)
        top_videos_csv = top_videos_df.to_csv(index=False)
        
        # Extract channel names
        channel_names = channel_data_df['Name'].tolist()
        
        top_videos_dict = top_videos_df.to_dict(orient='records')
        self.request.session['top_videos'] = top_videos_dict

        # Store in session
        self.request.session['channel_icons'] = channel_icons
        self.request.session['channel_names'] = channel_names
        self.request.session['channel_data_csv'] = channel_data_csv
        self.request.session['top_videos_csv'] = top_videos_csv
        self.request.session['durations'] = durations_list
        
        average_likes = channel_data_df['TotalLikes'].tolist()
        top_likes_channel = channel_data_df.sort_values('TotalLikes', ascending=False)['Name'].iloc[0]
        
        average_views = channel_data_df['TotalViews'].tolist()
        top_views_channel = channel_data_df.sort_values('TotalViews', ascending=False)['Name'].iloc[0]

        subs = channel_data_df['Subscriber count'].tolist()
        top_subs_channel = channel_data_df.sort_values('Subscriber count', ascending=False)['Name'].iloc[0]

        mostUsedCategories = channel_data_df['mostUsedCategories'].tolist()
        topTags = channel_data_df['Top tags'].tolist()
        
        self.request.session['mostUsedCategories'] = mostUsedCategories
        self.request.session['topTags'] = topTags
        self.request.session['average_likes'] = average_likes
        self.request.session['top_likes_channel'] = top_likes_channel
        self.request.session['average_views'] = average_views
        self.request.session['top_views_channel'] = top_views_channel
        self.request.session['subs'] = subs
        self.request.session['top_subs_channel'] = top_subs_channel
        self.request.session['type'] = entity_type
        self.request.session['durations'] = durations_list

        return HttpResponseRedirect(reverse('competitive_analysis_output')) 


# URL pattern would look something like this:
# path('analysis/', AnalysisWizard.as_view())
import json

def competitive_analysis_output_view(request):
    channel_icons = request.session.get('channel_icons', [])
    channel_names = request.session.get('channel_names', [])
    top_videos = request.session.get('top_videos', [])
    durations = request.session.get('durations', [])
    output_data = {
        'average_likes': request.session['average_likes'],
        'average_views': request.session['average_views'],
        'subs': request.session['subs'],
        'channel_names': channel_names,
        'durations': durations,
        'mostUsedCategories': request.session.get('mostUsedCategories', []),
        'topTags': request.session.get('topTags', []),
        
    }
    json_data = json.dumps(output_data)
    

    channels = zip(channel_icons, channel_names)
    channels_tags = zip(request.session.get('topTags', []), channel_names)
    channel_names = channel_names
    zipped_channels = zip(channel_icons, channel_names, top_videos)


    context = {
        'zipped_channels': zipped_channels,
        'channel_names': channel_names,
        'channel_icons': channel_icons,
        'channels': channels,
        'json_data': json_data,
        'top_likes_channel': request.session['top_likes_channel'],
        'top_views_channel': request.session['top_views_channel'],
        'top_subs_channel': request.session['top_subs_channel'],
        'type': request.session['type'],
        'top_videos': top_videos,
        'output_data': output_data,
        'channels_tags': channels_tags,
        'docx_file':'competitive_analysis.docx'
        
    }
    return render(request, 'features_pages/competitive_analysis/competitive_analysis_output.html', context)

def dataset_zipped_output(request):
    # Handle the output display here
    # Retrieve the CSV data from the session
    channel_data_csv = request.session.get('channel_data_csv', '')
    top_videos_csv = request.session.get('top_videos_csv', '')
    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        zip_file.writestr('channel_data.csv', channel_data_csv)
        zip_file.writestr('top_videos.csv', top_videos_csv)

    # Set up the HttpResponse
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="competitive_analysis_datasets.zip"'

    return response


@login_required
def video_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'features_pages/video_analysis/video_analysis_details.html')

from .forms import VideoAnalysisInputForm
from .models import VideoAnalysisHistory

class VideoAnalysisWizard(SessionWizardView):
    form_list = [VideoAnalysisInputForm]
    template_name = 'features_pages/video_analysis/video_analysis_forms.html'  
    
    def get_form_initial(self, step):
        initial = super().get_form_initial(step)
        history_id = self.kwargs.get('history_id')

        if history_id and step == '0':  # Assuming '0' is the step of VideoAnalysisInputForm
            history = get_object_or_404(VideoAnalysisHistory, id=history_id, user=self.request.user)
            initial.update({
                'video_url': history.video_url,
                # Add other fields here as necessary
            })
        return initial
    
    def done(self, form_list, **kwargs):
        # Process the cleaned data
        youtube = get_youtube_client()
        cleaned_data = self.get_all_cleaned_data()
        video_url = cleaned_data.get('video_url')
        openai_api_key = config('OPENAI_API_KEY')

        history_id = self.kwargs.get('history_id')
        if history_id:
            # Update the existing history record
            history = get_object_or_404(VideoAnalysisHistory, id=history_id, user=self.request.user)
            history.video_url = cleaned_data.get('video_url')
            # Update other fields as necessary
            history.save()
        else:
            VideoAnalysisHistory.objects.create(
            user=self.request.user,
            video_url=cleaned_data.get('video_url'),
            # Add other fields as necessary
            )
        
        video_id = extractIdFromUrl(video_url)
        video_info_df, comments_df, emotion_counts, top_comments_by_emotion = video_analysis(youtube, video_id)
        if 'engagementScore' in video_info_df.columns:
            video_info_df.drop('engagementScore', axis=1, inplace=True)
        video_info_csv = video_info_df.to_csv(index=False)
        comments_csv = comments_df.to_csv(index=False)
        
        self.request.session['video_info_csv'] = video_info_csv
        self.request.session['comments_csv'] = comments_csv
        
        
        self.request.session['video_info_dict'] = video_info_df.iloc[0].to_dict()
        self.request.session['emotion_counts'] = emotion_counts.to_dict()
        self.request.session['top_comments_by_emotion'] = top_comments_by_emotion
        
        # Define the output directory for audio files
        output_dir = os.path.join(settings.MEDIA_ROOT, 'audio')

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Download audio
        mp3_file = download_audio_from_youtube(video_url, output_dir)

        # Transcribe the audio file
        transcript = transcribe_youtube_video(mp3_file)

        # Summarize the transcript
        summary = summarize_youtube_video(transcript, openai_api_key)

        self.request.session['transcript'] = transcript
        self.request.session['summary'] = summary
        
        return HttpResponseRedirect(reverse('video_analysis_output'))  # Use the name of the URL pattern


def video_analysis_output_view(request):
    
    output_data = {
        "video_info_dict": request.session['video_info_dict'],
        'transcript': request.session['transcript'],
        'summary': request.session['summary']
    }
    
    if request.session['emotion_counts'] == {}:
        del request.session['emotion_counts']
    else:
        output_data["emotion_counts"] = request.session['emotion_counts']
    
    if request.session['top_comments_by_emotion'] == []:
        del request.session['top_comments_by_emotion']
    else:
        output_data["top_comments_by_emotion"] = request.session['top_comments_by_emotion']
     
    
    json_data = json.dumps(output_data)
    
    
    summary = request.session['summary']
    transcript = request.session['transcript']
    doc_output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
    
    docx_file_path = create_word_document(transcript, summary, "video_analysis", doc_output_dir)
    
    context= {'json_data': json_data, 
              'transcript': request.session['transcript'],
            'summary': request.session['summary'],
            "video_info_dict": request.session['video_info_dict'],
            'docx_file': docx_file_path,
            'top_comments_by_emotion': request.session['top_comments_by_emotion']
        }
    
    return render(request, 'features_pages/video_analysis/video_analysis_output.html', context)



import zipfile
import io

def dataset_zipped_output_video_analysis(request):
    # Handle the output display here
    # Retrieve the CSV data from the session
    video_info_csv = request.session.get('video_info_csv', '')
    comments_csv = request.session.get('comments_csv', '')
    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        zip_file.writestr('video_info_csv.csv', video_info_csv)
        zip_file.writestr('comments_csv.csv', comments_csv)

    # Set up the HttpResponse
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="video_analysis_datasets.zip"'

    return response

from django.http import FileResponse, Http404
import os

def download_docx(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, 'documents', filename)

    if not os.path.exists(file_path):
        raise Http404("File not found.")

    try:
        fh = open(file_path, 'rb')
        response = FileResponse(fh, as_attachment=True, filename=filename)
        return response
    except Exception as e:
        raise Http404(f"An error occurred: {str(e)}")
    
@login_required
def playlist_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'features_pages/playlist_analysis/playlist_analysis_details.html')

from .forms import PlaylistAnalysisInputForm
from .models import  PlaylistAnalysisHistory
class PlaylistAnalysisWizard(SessionWizardView):
    form_list = [PlaylistAnalysisInputForm]
    template_name = 'features_pages/playlist_analysis/playlist_analysis_forms.html'  
    
    def get_form_initial(self, step):
        initial = super().get_form_initial(step)
        history_id = self.kwargs.get('history_id')

        if history_id and step == '0':  # Assuming '0' is the step of PlaylistAnalysisInputForm
            history = get_object_or_404(PlaylistAnalysisHistory, id=history_id, user=self.request.user)
            initial.update({
                'playlist_url': history.playlist_url,
                # Add other fields as necessary
            })
        return initial
    
    def done(self, form_list, **kwargs):
        # Process the cleaned data
        youtube = get_youtube_client()
        cleaned_data = self.get_all_cleaned_data()
        plsylist_url = cleaned_data.get('playlist_url')
        
        history_id = self.kwargs.get('history_id')
        if history_id:
            # Update the existing history record
            history = get_object_or_404(PlaylistAnalysisHistory, id=history_id, user=self.request.user)
            history.playlist_url = cleaned_data.get('playlist_url')
            # Update other fields as necessary
            history.save()
        else:
            PlaylistAnalysisHistory.objects.create(
            user=self.request.user,
            playlist_url=cleaned_data.get('playlist_url'),
            # Add other fields as necessary
            )
        
        playlist_id = extractIdFromUrl(plsylist_url)
        playlist_info_df, all_videos_info_df, top_5_videos, worst_5_videos, top_5_comments_analysis, worst_5_comments_analysis = analyze_playlist(youtube, playlist_id)
        if 'engagementScore' in all_videos_info_df.columns:
            all_videos_info_df.drop('engagementScore', axis=1, inplace=True)
        playlist_info_csv = playlist_info_df.to_csv(index=False)
        all_videos_info_csv = all_videos_info_df.to_csv(index=False)
        
        self.request.session['playlist_info_csv'] = playlist_info_csv
        self.request.session['all_videos_info_csv'] = all_videos_info_csv
        
        self.request.session['title'] = playlist_info_df['title'].iloc[0]
        self.request.session['description'] = playlist_info_df['description'].iloc[0]
        self.request.session['publishedAt'] = playlist_info_df['publishedAt'].iloc[0]
        self.request.session['uniqueTags'] = playlist_info_df['uniqueTags'].iloc[0]
        self.request.session['thumbnail'] = playlist_info_df['thumbnail'].iloc[0]

       # Convert Pandas int64 values to native Python types for JSON serialization
        self.request.session['videoCount'] = int(playlist_info_df['videoCount'].iloc[0])
        self.request.session['totalViews'] = int(playlist_info_df['totalViews'].iloc[0])
        self.request.session['totalLikes'] = int(playlist_info_df['totalLikes'].iloc[0])
        self.request.session['totalComments'] = int(playlist_info_df['totalComments'].iloc[0])
        self.request.session['average_duration'] = float(playlist_info_df['average_duration'].iloc[0])

        # Convert the lists to native Python lists
        self.request.session['videos_publishedAt'] = all_videos_info_df['publishedAt'].tolist()
        self.request.session['videos_duration'] = all_videos_info_df['duration'].tolist()
        self.request.session['videos_views'] = all_videos_info_df['viewsCount'].tolist()
        self.request.session['videos_likes'] = all_videos_info_df['likesCount'].tolist()
        self.request.session['videos_commentCount'] = all_videos_info_df['commentCount'].tolist()
            
        self.request.session['top_5_videos'] = top_5_videos.to_dict(orient='records')
        self.request.session['worst_5_videos'] = worst_5_videos.to_dict(orient='records')
        if top_5_comments_analysis != []:
            top_5_comments_analysis_dist = top_5_comments_analysis[0].to_dict()
            top_5_comments = top_5_comments_analysis[1]
            self.request.session['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist
            self.request.session['top_5_comments'] = top_5_comments
        else:
            if 'top_5_comments_analysis_dist' in self.request.session:
                del self.request.session['top_5_comments_analysis_dist']
                del self.request.session['top_5_comments']
                
        if worst_5_comments_analysis != []:
            worst_5_comments_analysis_dist = worst_5_comments_analysis[0].to_dict()
            worst_5_comments = worst_5_comments_analysis[1]
            self.request.session['worst_5_comments_analysis_dist'] = worst_5_comments_analysis_dist
            self.request.session['worst_5_comments'] = worst_5_comments
        else:
            if 'worst_5_comments_analysis_dist' in self.request.session:
                del self.request.session['worst_5_comments_analysis_dist']
                del self.request.session['worst_5_comments']
                
        return HttpResponseRedirect(reverse('playlist_analysis_output'))  # Use the name of the URL pattern

import math
from datetime import datetime


def parse_datetime(date_str):
    for fmt in ('%Y-%m-%dT%H:%M:%S.%fZ', '%Y-%m-%dT%H:%M:%SZ'):
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            pass
    raise ValueError(f"time data {date_str} does not match expected formats")

def playlist_analysis_output_view(request):
    
    output_data = {
        'top_5_videos': request.session['top_5_videos'],
        'worst_5_videos': request.session['worst_5_videos'],
        'uniqueTags': request.session['uniqueTags'],
        'videos_publishedAt': request.session['videos_publishedAt'],
        'videos_duration': request.session['videos_duration'],
        'videos_likes': request.session['videos_likes'],
        'videos_views': request.session['videos_views'],
        'videos_commentCount': request.session['videos_commentCount']
    }
    
    if('top_5_comments_analysis_dist' in request.session): 
        top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
        top_5_comments = request.session['top_5_comments']
        output_data['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist
        output_data['top_5_comments'] = top_5_comments

    if('worst_5_comments_analysis_dist' in request.session): 
        worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
        worst_5_comments = request.session['worst_5_comments']
        output_data['worst_5_comments_analysis_dist'] = worst_5_comments_analysis_dist
        output_data['worst_5_comments'] = worst_5_comments

    
    publishedAt = request.session['publishedAt']
    
    date_obj = parse_datetime(publishedAt)

    # Format to 'YYYY MMM DD'
    formatted_date = date_obj.strftime('%Y %b %d')
    print(formatted_date)
    
    top_5_videos = request.session['top_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in top_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

    worst_5_videos = request.session['worst_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in worst_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

        
    json_data = json.dumps(output_data)
    
    context = {'json_data': json_data,
              'top_5_videos': top_5_videos,
              'worst_5_videos': worst_5_videos,
              'title': request.session['title'],
              'description': request.session['description'],
              'thumbnail': request.session['thumbnail'],
              'playlist_publishedAt': formatted_date,
              'videoCount': request.session['videoCount'],
              'totalViews': request.session['totalViews'],
              'totalLikes': request.session['totalLikes'],
              'totalComments': request.session['totalComments'],
              'average_duration': request.session['average_duration'],
              'uniqueTags': request.session['uniqueTags'],
              'videos_publishedAt': request.session['videos_publishedAt'],
              'videos_duration': request.session['videos_duration'],
              'videos_likes': request.session['videos_likes'],
              'videos_views': request.session['videos_views'],
              'videos_commentCount': request.session['videos_commentCount'],
              'docx_file': 'playlist_analysis.docx'
    }
    if('top_5_comments_analysis_dist' in request.session): 
        top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
        top_5_comments = request.session['top_5_comments']
        context['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist
        context['top_5_comments'] = top_5_comments

    if('worst_5_comments_analysis_dist' in request.session): 
        worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
        worst_5_comments = request.session['worst_5_comments']
        context['worst_5_comments_analysis_dist'] = worst_5_comments_analysis_dist
        context['worst_5_comments'] = worst_5_comments

    return render(request, 'features_pages/playlist_analysis/playlist_analysis_output.html', context)

def playlist_dataset_zipped_output(request):
    # Handle the output display here
    # Retrieve the CSV data from the session
    playlist_info_csv = request.session.get('playlist_info_csv', '')
    all_videos_info_csv = request.session.get('all_videos_info_csv', '')
    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        zip_file.writestr('playlist_info_csv.csv', playlist_info_csv)
        zip_file.writestr('all_videos_info_csv.csv', all_videos_info_csv)

    # Set up the HttpResponse
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="playlist_analysis_datasets.zip"'

    return response





@login_required
def channel_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'features_pages/channel_analysis/channel_analysis_details.html')

from .forms import ChannelAnalysisInputForm
from .models import ChannelAnalysisHistory

class ChannelAnalysisWizard(SessionWizardView):
    form_list = [ChannelAnalysisInputForm]
    template_name = 'features_pages/channel_analysis/channel_analysis_forms.html'  
    
    def get_form_initial(self, step):
        initial = super().get_form_initial(step)
        history_id = self.kwargs.get('history_id')

        if history_id and step == '0':  # Assuming '0' is the step of ChannelAnalysisInputForm
            history = get_object_or_404(ChannelAnalysisHistory, id=history_id, user=self.request.user)
            initial.update({
                'channel_url': history.channel_url,
                # Add other fields here as necessary
            })
        return initial
    
    def done(self, form_list, **kwargs):
        # Process the cleaned data
        youtube = get_youtube_client()
        cleaned_data = self.get_all_cleaned_data()
        channel_url = cleaned_data.get('channel_url')
        
        history_id = self.kwargs.get('history_id')
        if history_id:
            # Update the existing history record
            history = get_object_or_404(ChannelAnalysisHistory, id=history_id, user=self.request.user)
            history.channel_url = cleaned_data.get('channel_url')
            # Update other fields as necessary
            history.save()
        else:
            # Create a new history record
            ChannelAnalysisHistory.objects.create(
                user=self.request.user,
                channel_url=cleaned_data.get('channel_url'),
                # Add other fields as necessary
            )
        channel_id = extractIdFromUrl(channel_url)
        channel_df, all_videos_df, all_playlists_df, top_3_videos, worst_3_videos, top_3_comments_analysis, worst_3_comments_analysis = analyze_channel(youtube, channel_id)
        if 'engagementScore' in all_videos_df.columns:
            all_videos_df.drop('engagementScore', axis=1, inplace=True)
        channel_df_csv = channel_df.to_csv(index=False)
        all_videos_csv = all_videos_df.to_csv(index=False)
        all_playlists_csv = all_playlists_df.to_csv(index=False)
        
        self.request.session['channel_df_csv'] = channel_df_csv
        self.request.session['all_videos_info_csv'] = all_videos_csv
        self.request.session['all_playlists_dict'] = all_playlists_df.to_dict(orient='records')

        self.request.session['all_playlists_csv'] = all_playlists_csv

        self.request.session['title'] = channel_df['Channel Name'].iloc[0]
        self.request.session['description'] = channel_df['description'].iloc[0]
        self.request.session['publishedAt'] = channel_df['publishedAt'].iloc[0]
        self.request.session['uniqueTags'] = list(channel_df['uniqueTags'].iloc[0])
        self.request.session['thumbnail'] = channel_df['thumbnail'].iloc[0]
        self.request.session['mostUsedCategories'] = channel_df['mostUsedCategories'].iloc[0]
        self.request.session['average_duration'] = int(channel_df['average_duration'].iloc[0])

       # Convert Pandas int64 values to native Python types for JSON serialization
        self.request.session['videoCount'] = int(channel_df['videoCount'].iloc[0])
        self.request.session['totalViews'] = int(channel_df['viewCount'].iloc[0])
        self.request.session['totalLikes'] = int(channel_df['likesCount'].iloc[0])
        self.request.session['totalComments'] = int(channel_df['commentCount'].iloc[0])
        self.request.session['subscriberCount'] = float(channel_df['subscriberCount'].iloc[0])
        self.request.session['PlaylistCount'] = int(channel_df['PlaylistCount'].iloc[0])

        # Convert the lists to native Python lists
        self.request.session['videos_publishedAt'] = all_videos_df['publishedAt'].tolist()
        self.request.session['videos_duration'] = all_videos_df['duration'].tolist()
        self.request.session['videos_views'] = all_videos_df['viewsCount'].tolist()
        self.request.session['videos_likes'] = all_videos_df['likesCount'].tolist()
        self.request.session['videos_commentCount'] = all_videos_df['commentCount'].tolist()
            
        self.request.session['top_5_videos'] = top_3_videos.to_dict(orient='records')
        self.request.session['worst_5_videos'] = worst_3_videos.to_dict(orient='records')
        if top_3_comments_analysis != []:
            top_5_comments_analysis_dist = top_3_comments_analysis[0].to_dict()
            top_5_comments = top_3_comments_analysis[1]
            self.request.session['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist
            self.request.session['top_5_comments'] = top_5_comments
        else:
            if 'top_5_comments_analysis_dist' in self.request.session:
                del self.request.session['top_5_comments_analysis_dist']
                del self.request.session['top_5_comments']
        if worst_3_comments_analysis != []:
            worst_5_comments_analysis_dist = worst_3_comments_analysis[0].to_dict()
            worst_5_comments = worst_3_comments_analysis[1]
            self.request.session['worst_5_comments_analysis_dist'] = worst_5_comments_analysis_dist
            self.request.session['worst_5_comments'] = worst_5_comments
        else:
            if 'worst_5_comments_analysis_dist' in self.request.session:
                del self.request.session['worst_5_comments_analysis_dist']
                del self.request.session['worst_5_comments']
        
        return HttpResponseRedirect(reverse('channel_analysis_output'))  # Use the name of the URL pattern

import math
from datetime import datetime

def channel_analysis_output_view(request):
    
    output_data = {
        'top_5_videos': request.session['top_5_videos'],
        'worst_5_videos': request.session['worst_5_videos'],
        'uniqueTags': request.session['uniqueTags'],
        'videos_publishedAt': request.session['videos_publishedAt'],
        'videos_duration': request.session['videos_duration'],
        'videos_likes': request.session['videos_likes'],
        'videos_views': request.session['videos_views'],
        'videos_commentCount': request.session['videos_commentCount'],
        'all_playlists_dict': request.session['all_playlists_dict'],
        'videoCount': request.session['videoCount'],
        'totalViews': request.session['totalViews'],
        'totalLikes': request.session['totalLikes'],
        'totalComments': request.session['totalComments'],
        'mostUsedCategories': request.session['mostUsedCategories'],


    }
    
    if('top_5_comments_analysis_dist' in request.session): 
        top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
        top_5_comments = request.session['top_5_comments']
        output_data['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist
        output_data['top_5_comments'] = top_5_comments

    if('worst_5_comments_analysis_dist' in request.session): 
        worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
        worst_5_comments = request.session['worst_5_comments']
        output_data['worst_5_comments_analysis_dist'] = worst_5_comments_analysis_dist
        output_data['worst_5_comments'] = worst_5_comments
    
    publishedAt = request.session['publishedAt']
    
    date_obj = parse_datetime(publishedAt)

    # Format to 'YYYY MMM DD'
    formatted_date = date_obj.strftime('%Y %b %d')
    print(formatted_date)
    
    top_5_videos = request.session['top_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in top_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

    worst_5_videos = request.session['worst_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in worst_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

        
    json_data = json.dumps(output_data)
    
    context= {'json_data': json_data,
              'top_5_videos': top_5_videos,
              'worst_5_videos': worst_5_videos,
              'title': request.session['title'],
              'description': request.session['description'],
              'thumbnail': request.session['thumbnail'],
              'playlist_publishedAt': formatted_date,
              'videoCount': request.session['videoCount'],
              'totalViews': request.session['totalViews'],
              'totalLikes': request.session['totalLikes'],
              'totalComments': request.session['totalComments'],
              'average_duration': request.session['average_duration'],
              'uniqueTags': request.session['uniqueTags'],
              'videos_publishedAt': request.session['videos_publishedAt'],
              'videos_duration': request.session['videos_duration'],
              'videos_likes': request.session['videos_likes'],
              'videos_views': request.session['videos_views'],
              'videos_commentCount': request.session['videos_commentCount'],
              'docx_file': 'channel_analysis.docx',
              }
    
    if('top_5_comments_analysis_dist' in request.session): 
        top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
        top_5_comments = request.session['top_5_comments']
        context['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist
        context['top_5_comments'] = top_5_comments

    if('worst_5_comments_analysis_dist' in request.session): 
        worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
        worst_5_comments = request.session['worst_5_comments']
        context['worst_5_comments_analysis_dist'] = worst_5_comments_analysis_dist
        context['worst_5_comments'] = worst_5_comments
    
    return render(request, 'features_pages/channel_analysis/channel_analysis_output.html', context)

def channel_dataset_zipped_output(request):
    # Handle the output display here
    # Retrieve the CSV data from the session
    channel_df_csv = request.session.get('channel_df_csv', '')
    all_videos_info_csv = request.session.get('all_videos_info_csv', '')
    all_playlists_csv = request.session.get('all_playlists_csv', '')
    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        zip_file.writestr('channel_df_csv.csv', channel_df_csv)
        zip_file.writestr('all_videos_info_csv.csv', all_videos_info_csv)
        zip_file.writestr('all_playlists_csv.csv', all_playlists_csv)


    # Set up the HttpResponse
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="channel_analysis_datasets.zip"'

    return response


@login_required
def topic_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'features_pages/topic_analysis/topic_analysis_details.html')


def topictrend_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'instafeatures_pages/topictrend_analysis/topictrend_analysis_details.html')

def posts_analysis_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'instafeatures_pages/posts_analysis/posts_analysis_details.html')

def engagement_history_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'instafeatures_pages/engagement_history/engagement_history_details.html')

def people_analytics_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'instafeatures_pages/people_analytics/people_analytics_details.html')

def comparative_study_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'instafeatures_pages/comparative_study/comparative_study_details.html')

def instagram_reporting_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'instafeatures_pages/instagram_reporting/instagram_reporting_details.html')


from .forms import ChannelAnalysisInputForm

class TopicAnalysisWizard(SessionWizardView):
    form_list = [YouTubeSearchForm]
    template_name = 'features_pages/topic_analysis/topic_analysis_forms.html'

    def get_form_initial(self, step):
        initial = super().get_form_initial(step)
        history_id = self.kwargs.get('history_id')

        if history_id and step == '0':  # Assuming '0' is the step you want to pre-fill
            history = get_object_or_404(TopicAnalysisHistory, id=history_id)
            initial.update({
                'search_query': history.search_query,
                'order': history.order,
                'region_code': history.region_code,
                'language': history.language,
                # Add other fields as necessary
            })
        return initial

    
    def done(self, form_list, **kwargs):
        # Process the cleaned data
        youtube = get_youtube_client()
        cleaned_data = self.get_all_cleaned_data()
        query = cleaned_data.get('search_query')
        order = cleaned_data.get('order')
        region_code = cleaned_data.get('region_code')
        language = cleaned_data.get('language')
        
        videos_df, channels_df, top_5_videos, comments_df, top_5_comments_analysis, keybert_keywords = topic_analysis(youtube, query, order, region_code, language, max_results=100)
        if 'engagementScore' in videos_df.columns:
            videos_df.drop('engagementScore', axis=1, inplace=True)
        channel_df_csv = channels_df.to_csv(index=False)
        all_videos_csv = videos_df.to_csv(index=False)
        comments_csv = comments_df.to_csv(index=False)
        
        self.request.session['channel_df_csv'] = channel_df_csv
        self.request.session['all_videos_info_csv'] = all_videos_csv
        self.request.session['comments_csv'] = comments_csv

        self.request.session['videos_dict'] = videos_df.to_dict(orient='records')
        self.request.session['channels_dict'] = channels_df.to_dict(orient='records')
        self.request.session['top_5_videos'] = top_5_videos.to_dict(orient='records')
        self.request.session['comments_dict'] = comments_df.to_dict(orient='records')
        top_5_comments_analysis_dist = top_5_comments_analysis[0]
        top_5_comments = top_5_comments_analysis[1]
        
        self.request.session['top_5_comments_analysis_dist'] = top_5_comments_analysis_dist.to_dict()
        self.request.session['top_5_comments'] = top_5_comments
        self.request.session['keybert_keywords'] = keybert_keywords


        TopicAnalysisHistory.objects.create(
            user=self.request.user,
            search_query=cleaned_data.get('search_query'),
            order=cleaned_data.get('order'),
            region_code = cleaned_data.get('region_code'),
            language = cleaned_data.get('language')
        )
        return HttpResponseRedirect(reverse('topic_analysis_output'))  # Use the name of the URL pattern

import math
from datetime import datetime

def topic_analysis_output_view(request):

    
    output_data = {
        'top_5_videos': request.session['top_5_videos'],
        'channels_dict': request.session['channels_dict'],
        'videos_dict': request.session['videos_dict'],
        'comments_dict': request.session['comments_dict'],
        'top_5_comments_analysis_dist': request.session['top_5_comments_analysis_dist'],
        'top_5_comments': request.session['top_5_comments'],
        'keybert_keywords': request.session['keybert_keywords'],
    }
    
    
    top_5_videos = request.session['top_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in top_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

        
    json_data = json.dumps(output_data)
    
    context= {
        'json_data': json_data,
        'top_5_videos': top_5_videos,
        'channels_dict': request.session['channels_dict'],
        'videos_dict': request.session['videos_dict'],
        'comments_dict': request.session['comments_dict'],
        'top_5_comments_analysis_dist': request.session['top_5_comments_analysis_dist'],
        'top_5_comments': request.session['top_5_comments'],
        'docx_file': 'topic_analysis.docx',
              }
    
    
    return render(request, 'features_pages/topic_analysis/topic_analysis_output.html', context)

def topic_dataset_zipped_output(request):
    # Handle the output display here
    # Retrieve the CSV data from the session
    channel_df_csv = request.session.get('channel_df_csv', '')
    all_videos_info_csv = request.session.get('all_videos_info_csv', '')
    comments_csv = request.session.get('comments_csv', '')
    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        zip_file.writestr('channel_df_csv.csv', channel_df_csv)
        zip_file.writestr('all_videos_info_csv.csv', all_videos_info_csv)
        zip_file.writestr('comments_csv.csv', comments_csv)


    # Set up the HttpResponse
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="topic_analysis_datasets.zip"'

    return response

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import base64
from docx import Document
from io import BytesIO
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@csrf_exempt
def doc_competitive(request):
    channel_icons = request.session.get('channel_icons', [])
    channel_names = request.session.get('channel_names', [])
    top_videos = request.session.get('top_videos', [])
    durations = request.session.get('durations', [])
    output_data = {
        'average_likes': request.session['average_likes'],
        'average_views': request.session['average_views'],
        'subs': request.session['subs'],
        'channel_names': channel_names,
        'durations': durations,
        'mostUsedCategories': request.session.get('mostUsedCategories', []),
        'topTags': request.session.get('topTags', []),
    }
    json_data = json.dumps(output_data)
    

    # Zip the lists together in the view
    channels = zip(channel_icons, channel_names)
    channels_tags = zip(request.session.get('topTags', []), channel_names)
    context = {
        'channels': channels,
        'json_data': json_data,
        'top_likes_channel': request.session['top_likes_channel'],
        'top_views_channel': request.session['top_views_channel'],
        'top_subs_channel': request.session['top_subs_channel'],
        'type': request.session['type'],
        'top_videos': top_videos,
        'output_data': output_data,
        'channels_tags': channels_tags
        
    }
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Compeitive Analysis', 0)
        doc.add_paragraph('Our youtube competitive analysis will provide you with customizable dataset, statistics, graphs and interpretaions to make your work with data easier.')
        
        doc.add_heading('Analysed playlist:', level=1)
        for name in channel_names:
            doc.add_paragraph(name, style='ListBullet')

        doc.add_heading('Top videos from each channel:', level=1)
        for video in top_videos:
            # Heading for each video title
            doc.add_heading(video['title'], level=2)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

            # Comments as bullet points
            if video['topComments']:  # Check if there are comments
                comments_para = doc.add_paragraph()
                comments_para.add_run('Top Comments:').bold = True
                for comment in video['topComments']:
                    doc.add_paragraph(comment, style='ListBullet')
                    
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "competitive_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)


@csrf_exempt
def doc_channel(request):
    
    channel_data= {
        
    
    'top_5_videos': request.session['top_5_videos'],
    'worst_5_videos': request.session['worst_5_videos'],
    'uniqueTags': request.session['uniqueTags'],
    'all_playlists_dict': request.session['all_playlists_dict'],         
    'title': request.session['title'],
    'description': request.session['description'],
    'thumbnail': request.session['thumbnail'],
    'videoCount': request.session['videoCount'],
    'totalViews': request.session['totalViews'],
    'totalLikes': request.session['totalLikes'],
    'totalComments': request.session['totalComments'],
    'average_duration': request.session['average_duration'],
    }
    
    
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Channel Analysis', 0)
        doc.add_paragraph('Our channel analysis will give you an overview over the channel, what does influence its performance, and a closer look on its top and worst performing videos...')
        
        doc.add_heading('Get An Overview', level=1)
        doc.add_heading(channel_data['title'], level=2)
        
        doc.add_heading('Discreption', level=3)
        doc.add_paragraph(channel_data['description'])
        doc.add_heading('Word tags', level=3)
        tags_str = ', '.join(channel_data['uniqueTags'])
        doc.add_paragraph(tags_str)
        doc.add_heading('Statistics', level=3)
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Video Count: ').bold = True
        stats_para.add_run(str(channel_data['videoCount']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total Views: ').bold = True
        stats_para.add_run(str(channel_data['totalViews']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total likes: ').bold = True
        stats_para.add_run(str(channel_data['totalLikes']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Comments Count: ').bold = True
        stats_para.add_run(str(channel_data['totalComments']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Videos Average Duration: ').bold = True
        stats_para.add_run(str(channel_data['average_duration']))  # Convert integer to string

        doc.add_heading('Top and worst videos', level=1)
        
        doc.add_heading('Top videos', level=2)
        doc.add_heading('Top videos info', level=3)
        for video in channel_data['top_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")
        
        
        if('top_5_comments_analysis_dist' in request.session): 
            top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
            top_5_comments = request.session['top_5_comments']

            doc.add_heading('Top videos Comments and Sentiment', level=3)
            for emotion, comment in top_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)
            
            
            
        doc.add_heading('Worst videos', level=2)
        doc.add_heading('Worst videos info', level=3)
        for video in channel_data['worst_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

        if('worst_5_comments_analysis_dist' in request.session): 
            worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
            worst_5_comments = request.session['worst_5_comments']
        
            doc.add_heading('Worst videos Comments and Sentiment', level=3)
            for emotion, comment in worst_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)

        
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "channel_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)



@csrf_exempt
def doc_playlist(request):
    
    channel_data= {
        
    
    'top_5_videos': request.session['top_5_videos'],
    'worst_5_videos': request.session['worst_5_videos'],
    'uniqueTags': request.session['uniqueTags'],
    'all_playlists_dict': request.session['all_playlists_dict'],         
    'title': request.session['title'],
    'description': request.session['description'],
    'thumbnail': request.session['thumbnail'],
    'videoCount': request.session['videoCount'],
    'totalViews': request.session['totalViews'],
    'totalLikes': request.session['totalLikes'],
    'totalComments': request.session['totalComments'],
    'average_duration': request.session['average_duration'],
    }
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Playlist Analysis', 0)
        doc.add_paragraph('Our playlist analysis will give you an overview over the playlist, what does influence its performance, and a closer look on its top and worst performing videos...')
        
        doc.add_heading('Get An Overview', level=1)
        doc.add_heading(channel_data['title'], level=2)
        
        doc.add_heading('Discreption', level=3)
        doc.add_paragraph(channel_data['description'])
        doc.add_heading('Word tags', level=3)
        tags_str = ', '.join(channel_data['uniqueTags'])
        doc.add_paragraph(tags_str)
        doc.add_heading('Statistics', level=3)
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Video Count: ').bold = True
        stats_para.add_run(str(channel_data['videoCount']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total Views: ').bold = True
        stats_para.add_run(str(channel_data['totalViews']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Total likes: ').bold = True
        stats_para.add_run(str(channel_data['totalLikes']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Comments Count: ').bold = True
        stats_para.add_run(str(channel_data['totalComments']))  # Convert integer to string
        stats_para = doc.add_paragraph(style='ListBullet')
        stats_para.add_run('Videos Average Duration: ').bold = True
        stats_para.add_run(str(channel_data['average_duration']))  # Convert integer to string

        doc.add_heading('Top and worst videos', level=1)
        
        doc.add_heading('Top videos', level=2)
        doc.add_heading('Top videos info', level=3)
        for video in channel_data['top_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

        if('top_5_comments_analysis_dist' in request.session): 
            top_5_comments_analysis_dist = request.session['top_5_comments_analysis_dist']
            top_5_comments = request.session['top_5_comments']

            doc.add_heading('Top videos Comments and Sentiment', level=3)
            for emotion, comment in top_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)

            
            
            
        doc.add_heading('Worst videos', level=2)
        doc.add_heading('Worst videos info', level=3)
        for video in channel_data['worst_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=4)

            # Bold description label
            desc_para = doc.add_paragraph()
            desc_para.add_run('Description: ').bold = True
            desc_para.add_run(video['description'])
            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")

        if('worst_5_comments_analysis_dist' in request.session): 
            worst_5_comments_analysis_dist = request.session['worst_5_comments_analysis_dist']
            worst_5_comments = request.session['worst_5_comments']
        
            doc.add_heading('Worst videos Comments and Sentiment', level=3)
            for emotion, comment in worst_5_comments.items():
                stats_para = doc.add_paragraph(style='ListBullet')
                stats_para.add_run(f"{emotion} :").bold = True
                stats_para.add_run(comment)
        
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "playlist_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)



@csrf_exempt
def doc_topic(request):
    
    top_5_videos = request.session['top_5_videos']
    # Assuming 'top_5_videos' is a list of dictionaries containing video information
    for video in top_5_videos:
        # Parse the 'publishedAt' date string to a datetime object
        date_obj = parse_datetime(video['publishedAt'])

        # Reformat the date to 'YYYY MMM DD' and update the video dictionary
        video['publishedAt'] = date_obj.strftime('%Y %b %d')

            
    context= {
        'top_5_videos': top_5_videos,
        'channels_dict': request.session['channels_dict'],
        'videos_dict': request.session['videos_dict'],
        'top_5_comments': request.session['top_5_comments'],
              }
              
    
    if request.method == 'POST':
        doc = Document()
        doc.add_heading('Topic Analysis', 0)
        doc.add_paragraph('Check out the metrics and insights to dive into the trends and discussions around your chosen subject.')
                
        
        doc.add_heading('Top channels', level=1)
        for channel in context['channels_dict']:
            # Heading for each video title
            doc.add_heading(channel['Name'], level=2)

            

            doc.add_heading('Statistics', level=3)
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Video Count: ').bold = True
            stats_para.add_run(str(channel['Video count']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Views average: ').bold = True
            stats_para.add_run(str(channel['TotalViews']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('likes average: ').bold = True
            stats_para.add_run(str(channel['TotalLikes']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Subscriber count: ').bold = True
            stats_para.add_run(str(channel['Subscriber count']))  # Convert integer to string
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run('Playlist count: ').bold = True
            stats_para.add_run(str(channel['Playlist count']))  # Convert integer to string
            
            # Bold description label
            desc_para = doc.add_paragraph()
            
            
        
        doc.add_heading('Top videos', level=1)
        for video in context['top_5_videos']:
            # Heading for each video title
            doc.add_heading(video['title'], level=2)

            
            # Bold statistics label
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['viewsCount']} views, {video['likesCount']} likes, {video['duration']} minutes")
            # Bold description label
            desc_para = doc.add_paragraph()
            
            
            
        doc.add_heading('Top videos Comments and Sentiment', level=1)
        for emotion, comment in context['top_5_comments'].items():
            stats_para = doc.add_paragraph(style='ListBullet')
            stats_para.add_run(f"{emotion} :").bold = True
            stats_para.add_run(comment)
            
            
    
        

        
        doc.add_heading('Get more insights with graphs:', level=1)
        
        data = json.loads(request.body)
        # The variable here should match what's sent from the frontend
        imgs_data = data['imgData']  # This should match the key in the JSON sent from the frontend
        for img_data in imgs_data:
            # Decode the base64 image
            img_data = base64.b64decode(img_data.split(',')[1])
            image_stream = BytesIO(img_data)
            # Add the image to the Word document
            doc.add_picture(image_stream)

        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "topic_analysis.docx")
        doc.save(file_path)

        return JsonResponse({'message': 'Document created successfully'})

    return JsonResponse({'error': 'Invalid request'}, status=400)
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt
from docx.text.run import Run
import docx
def add_hyperlink(paragraph, url, text, color, underline, heading=False):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    if not underline:
        u = OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    if heading:
        run = Run(new_run, paragraph)
        run.font.size = Pt(16)  # Adjust the size as per your heading style
        run.font.bold = True

    return hyperlink


def def_retrive(request):
              
        doc = Document()
        doc.add_heading('Videos Retriving', 0)
        doc.add_paragraph('Explore the collection below to discover content related to your search. Use these insights to enhance your understanding, create content, or simply enjoy the diversity of videos available on your topic of interest.')
                
        related_videos_full_dict = request.session['related_videos_full_dict']
        
        doc.add_heading('List of Videos', level=1)
        for video in related_videos_full_dict:
            p = doc.add_paragraph()
            add_hyperlink(p, video['URL'], video['Title'], '0000FF', False, heading=True)
            stats_para = doc.add_paragraph()
            stats_para.add_run('Channel: ').bold = True
            stats_para.add_run(f"{video['Channel']}")
            stats_para = doc.add_paragraph()
            stats_para.add_run('Statistics: ').bold = True
            stats_para.add_run(f"{video['Views']} views, {video['Likes']} likes, {video['Comments']} comments, {video['Duration']} minutes")
            stats_para = doc.add_paragraph()
            stats_para.add_run('Category: ').bold = True
            stats_para.add_run(f"{video['Category']}")
            stats_para = doc.add_paragraph()
            stats_para.add_run('Top Tags: ').bold = True
            stats_para.add_run(f"{video['Tags'][:5]}")
    
        # Save the document
        output_dir = os.path.join(settings.MEDIA_ROOT, 'documents')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        file_path = os.path.join(output_dir, "video_retriving.docx")
        doc.save(file_path)



@login_required
def video_retriving_details(request):
    # You can add code here to fetch and process inquiries
    return render(request, 'features_pages/video_retriving/video_retriving_details.html')

from .forms import VideoRetrivingInputForm
from .models import VideoRetrievingHistory
class VideoRetrivingWizard(SessionWizardView):
    form_list = [VideoRetrivingInputForm]
    template_name = 'features_pages/video_retriving/video_retriving_forms.html'  
    
    def get_form_initial(self, step):
        initial = super().get_form_initial(step)
        history_id = self.kwargs.get('history_id')

        if history_id and step == '0':  # Assuming '0' is the step of VideoRetrivingInputForm
            history = get_object_or_404(VideoRetrievingHistory, id=history_id, user=self.request.user)
            initial.update({
                'search_query': history.search_query,
                'order': history.order,
                'region_code': history.region_code,
                'language': history.language,
                'num_of_videos': history.num_of_videos,
                # Add other fields as necessary
            })
        return initial
    
    def done(self, form_list, **kwargs):
        # Process the cleaned data
        youtube = get_youtube_client()
        cleaned_data = self.get_all_cleaned_data()
        video_url = cleaned_data.get('search_query')
        order = cleaned_data.get('order')
        region_code = cleaned_data.get('region_code')
        language = cleaned_data.get('language')
        num_of_videos = cleaned_data.get('num_of_videos')
        
        history_id = self.kwargs.get('history_id')
        if history_id:
            # Update the existing history record
            history = get_object_or_404(VideoRetrievingHistory, id=history_id, user=self.request.user)
            history.search_query=cleaned_data.get('search_query'),
            history.order=cleaned_data.get('order'),
            history.region_code=cleaned_data.get('region_code'),
            history.language=cleaned_data.get('language'),
            history.num_of_videos=cleaned_data.get('num_of_videos'),
            # Update other fields as necessary
            history.save()
        else:
            VideoRetrievingHistory.objects.create(
            user=self.request.user,
            search_query=cleaned_data.get('search_query'),
            order=cleaned_data.get('order'),
            region_code=cleaned_data.get('region_code'),
            language=cleaned_data.get('language'),
            num_of_videos=cleaned_data.get('num_of_videos'),
            )
        
        video_id = extractIdFromUrl(video_url)
        
        related_videos_df = get_realted_videos(youtube, video_id, order, region_code, language, num_of_videos)
        if 'engagementScore' in related_videos_df.columns:
            related_videos_df.drop('engagementScore', axis=1, inplace=True)
            
        related_videos_csv = related_videos_df.to_csv(index=False)
        related_videos_dict = related_videos_df.to_dict(orient='records')[:8]
        related_videos_full_dict = related_videos_df.to_dict(orient='records')
        self.request.session['related_videos_csv'] = related_videos_csv
        self.request.session['related_videos_dict'] = related_videos_dict
        self.request.session['related_videos_full_dict'] = related_videos_full_dict
        return HttpResponseRedirect(reverse('video_retriving_output'))  # Use the name of the URL pattern


def video_retriving_output_view(request):
    def_retrive(request)
    
    context = {
        'related_videos_csv': request.session['related_videos_csv'],
        'related_videos_dict':request.session['related_videos_dict'],
        'related_videos_full_dict': request.session['related_videos_full_dict'],
        'docx_file': 'video_retriving.docx',
    }
    return render(request, 'features_pages/video_retriving/video_retriving_output.html', context)



import zipfile
import io

def dataset_zipped_output_retriving(request):
    
    related_videos_csv = request.session.get('related_videos_csv', '')

    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
        zip_file.writestr('related_videos_csv.csv', related_videos_csv)

    # Set up the HttpResponse
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="related_videos_dataset.zip"'

    return response


from django.shortcuts import render, get_object_or_404
from .models import TopicAnalysisHistory  # Import your model

def topic_analysis_detail(request, history_id):
    # Fetch the history instance or return 404 if not found
    history = get_object_or_404(TopicAnalysisHistory, pk=history_id, user=request.user)

    # Render the detail in a template
    return render(request, 'features_pages/topic_analysis/topic_analysis_detail.html', {'history': history})

def video_analysis_detail(request, history_id):
    history = get_object_or_404(VideoAnalysisHistory, pk=history_id, user=request.user)
    return render(request, 'features_pages/video_analysis/video_analysis_detail.html', {'history': history})

def playlist_analysis_detail(request, history_id):
    history = get_object_or_404(PlaylistAnalysisHistory, pk=history_id, user=request.user)
    return render(request, 'features_pages/playlist_analysis/playlist_analysis_detail.html', {'history': history})

def channel_analysis_detail(request, history_id):
    history = get_object_or_404(ChannelAnalysisHistory, pk=history_id, user=request.user)
    return render(request, 'features_pages/channel_analysis/channel_analysis_detail.html', {'history': history})

def video_retriving_detail(request, history_id):
    history = get_object_or_404(VideoRetrievingHistory, pk=history_id, user=request.user)
    return render(request, 'features_pages/video_retrieving/video_retrieving_detail.html', {'history': history})

def competitive_analysis_detail(request, history_id):
    history = get_object_or_404(CompetitiveAnalysisHistory, pk=history_id, user=request.user)
    return render(request, 'features_pages/competitive_analysis/competitive_analysis_detail.html', {'history': history})

from django.shortcuts import redirect, get_object_or_404

def get_model_by_type(history_type):
    if history_type == 'video':
        return VideoAnalysisHistory
    elif history_type == 'topic':
        return TopicAnalysisHistory
    elif history_type == 'playlist':
        return PlaylistAnalysisHistory
    elif history_type == 'channel':
        return ChannelAnalysisHistory
    elif history_type == 'video_retrieving':
        return VideoRetrievingHistory
    elif history_type == 'competitive':
        return CompetitiveAnalysisHistory
    else:
        raise ValueError("Unknown history type")

def delete_history(request, history_type, history_id):
    model = get_model_by_type(history_type)
    history = get_object_or_404(model, pk=history_id, user=request.user)
    history.delete()
    return redirect('base')  # Redirect to an appropriate page


# views.py
import traceback

from django.shortcuts import render
from django.http import JsonResponse
import openai
import json

def chat_view(request):
    
    function = {
    "name": "generate_code",
    "description": "Generates Python code for data manipulation based on user input, it takes in json string of the dataframe structure and the first 5 rows, and generate code to edit it based on user request, the end results should be a pandas dataframe named df.",
    "parameters": {
        "type": "object",
        "properties": {
        "generated_code": {
            "type": "string",
            "description": "Python code to be generated for manipulating a dataset. Format example: {'generated_code': 'all_videos_info_df= pd.DataFrame([all_videos_info_dict]) all_videos_info_df[\\'sum_likes_comments\\'] = all_videos_info_df[\\'likesCount\\'] + all_videos_info_df[\\'commentCount\\']'}"
        },
        "explanation": {
            "type": "string",
            "description": "Explain what has changed to ther user."
        },
        },
        "required": [
        "generated_code", "explanation"
        ]
    }}
    
    # Initialize or get the existing conversation from the session
    if 'conversation' not in request.session or request.method == 'GET':
        request.session['conversation'] = [
            {"role": "system", "content": "You are a Python code generator for data manipulation. Please provide your data manipulation request."},
            # Assuming the dataset JSON structure is passed initially from the view
            {"role": "assistant", "content": "Here is your dataset: [Your dataset structure here]"},
        ]
        
    if 'error_count' not in request.session:
        request.session['error_count'] = 0

    if request.method == 'POST':
        did_not_execute = True
        while did_not_execute:
            try:
                user_input = request.POST.get('user_input')

                request.session['conversation'].append({"role": "user", "content": user_input})

                response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo-0613",
                        messages=request.session['conversation'],
                        functions=[function],
                        function_call={
                            "name": "generate_code",
                            "data": {
                                "json_dataset": json.dumps([request.session['videos_dict'][:5]]),
                                "user_request": user_input
                            }
                        }
                    )

                # Process the response
                content = response.choices[0]["message"]["function_call"]["arguments"]
                content_json = json.loads(content)
                generated_code = content_json['generated_code']
                print(generated_code)
                # Attempt to execute the generated code
                df = pd.DataFrame(request.session['videos_dict'])
                exec(generated_code)

                # Reset error count on successful execution
                request.session['error_count'] = 0

                # Append successful response to the conversation
                response_content = generated_code

                # Ask if the user has more changes after successful execution
                request.session['conversation'].append({"role": "assistant", "content": "Do you have any more changes?"})

                # Send response back to user
                did_not_execute = False
                send_response = True

            except Exception as e:
                # print(str(e))
                # Increment error count
                request.session['error_count'] += 1

                full_traceback = traceback.format_exc().splitlines()
                error_traceback = "\n".join(full_traceback[-3:]) if len(full_traceback) >= 3 else str(e)
                print(error_traceback)
                # Append error message to the conversation for the assistant to process internally
                error_message = f"Error occurred, please try to modify, stick to the columns names in the example provided {json.dumps([request.session['videos_dict'][:1]])}, the error: {error_traceback}."
                request.session['conversation'].append({"role": "user", "content": error_message})
                
                # Check if errors have occurred 3 times consecutively
                if request.session['error_count'] >= 3:
                    # Inform the user about persistent issues
                    response_content = "Sorry, this request doesn't seem to work. Let's try something else."
                    request.session['error_count'] = 0  # Reset error count
                    send_response = True
                    did_not_execute = False
                else:
                    # Continue trying to resolve internally, do not send response to user yet
                    send_response = False
                    did_not_execute = True

        # Save the updated conversation and error count to the session
        request.session.modified = True

        if send_response:
            # Return the response
            return JsonResponse({"ai_response": response_content})

    return render(request, 'chat_template.html')

#edit the published at date to this format MM/YYYY

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

@csrf_exempt
def delete_selected_templates(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        selected_items = data.get('selectedItems', [])

        for item in selected_items:
            model = get_model_by_type(item['type'])
            model.objects.filter(id=item['id']).delete()

        return JsonResponse({'status': 'success', 'message': 'Templates deleted successfully.'})
    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)
